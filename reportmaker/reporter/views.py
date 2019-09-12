import mimetypes
import os

from django.http import HttpResponseRedirect, HttpResponse
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth import authenticate, login, logout
from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from django.views import generic

from .models import Report, Process
from .forms import ReportForm


class ReportListView(LoginRequiredMixin, generic.ListView):
	login_url = '/login'
	template_name = 'reporter/report_list.html'
	filters = {
		'id': '',
		'title': '',
		'title_match': '',
		'content': '',
		'content_match': '',
		'date': ''
	}

	def get_queryset(self):
		reports = Report.objects.all()
		if self.request.method == 'GET':
			self.filters['id'] = self.request.GET.get('id_filter')
			self.filters['title'] = self.request.GET.get('title_filter')
			self.filters['title_match'] = self.request.GET.get('title_match')
			self.filters['content'] = self.request.GET.get('content_filter')
			self.filters['content_match'] = self.request.GET.get('content_match')

			if self.filters['id']:
				reports = reports.filter(id=self.filters['id'])

			if self.filters['title']:
				if self.filters['title_match'] == 'exact':
					reports = reports.filter(title_text__iexact=self.filters['title'])
				elif self.filters['title_match'] == 'starts':
					reports = reports.filter(title_text__istartswith=self.filters['title'])
				elif self.filters['title_match'] == 'contains':
					reports = reports.filter(title_text__icontains=self.filters['title'])
				elif self.filters['title_match'] == 'ends':
					reports = reports.filter(title_text__iendswith=self.filters['title'])

			if self.filters['content']:
				if self.filters['content_match'] == 'exact':
					reports = reports.filter(content_text__iexact=self.filters['content'])
				elif self.filters['content_match'] == 'starts':
					reports = reports.filter(content_text__istartswith=self.filters['content'])
				elif self.filters['content_match'] == 'contains':
					reports = reports.filter(content_text__icontains=self.filters['content'])
				elif self.filters['content_match'] == 'ends':
					reports = reports.filter(content_text__iendswith=self.filters['content'])

		# TODO: filter by date

		return reports

	def get_context_data(self, **kwargs):
		context = super(ReportListView, self).get_context_data(**kwargs)
		context['filters'] = self.filters
		return context


class ProcessListView(LoginRequiredMixin, generic.ListView):
	login_url = '/login'
	template_name = 'reporter/report.html'

	filters = {
		'id': '',
		'name': '',
		'name_match': '',
		'state': '',
		'state_match': '',
		'ram': '',
		'date': ''
	}

	def get_queryset(self):
		processes = Process.objects.all()
		processes = processes.filter(report=self.kwargs.get('report_id'))

		if self.request.method == 'GET':
			self.filters['id'] = self.request.GET.get('id_filter')
			self.filters['name'] = self.request.GET.get('name_filter')
			self.filters['name_match'] = self.request.GET.get('name_match')
			self.filters['state'] = self.request.GET.get('state_filter')
			self.filters['state_match'] = self.request.GET.get('state_match')
			self.filters['ram'] = self.request.GET.get('ram_filter')
			self.filters['date'] = self.request.GET.get('date_filter')

			if self.filters['id']:
				processes = processes.filter(id=self.filters['id'])

			if self.filters['name']:
				if self.filters['name_match'] == 'exact':
					processes = processes.filter(name_text__iexact=self.filters['name'])
				elif self.filters['name_match'] == 'starts':
					processes = processes.filter(name_text__istartswith=self.filters['name'])
				elif self.filters['name_match'] == 'contains':
					processes = processes.filter(name_text__icontains=self.filters['name'])
				elif self.filters['name_match'] == 'ends':
					processes = processes.filter(name_text__iendswith=self.filters['name'])

			if self.filters['state']:
				if self.filters['state_match'] == 'exact':
					processes = processes.filter(state_text__iexact=self.filters['state'])
				elif self.filters['state_match'] == 'starts':
					processes = processes.filter(state_text__istartswith=self.filters['state'])
				elif self.filters['state_match'] == 'contains':
					processes = processes.filter(state_text__icontains=self.filters['state'])
				elif self.filters['state_match'] == 'ends':
					processes = processes.filter(state_text__iendswith=self.filters['state'])

			if self.filters['ram']:
				processes = processes.filter(ram_in_kb_int__startswith=int(self.filters['ram']))


		# TODO: filter by date

		return processes

	def get_context_data(self, **kwargs):
		context = super(ProcessListView, self).get_context_data(**kwargs)
		context['filters'] = self.filters
		context['report'] = get_object_or_404(Report, id=self.kwargs.get('report_id'))
		return context


def index(request):
	return render(request, 'reporter/index.html')


@login_required()
def report(request, report_id):
	report = get_object_or_404(Report, id=report_id)
	return render(request, 'reporter/report.html', {'report': report})


@login_required()
def remove(request, report_id):
	if request.method == 'POST':
		Report.objects.get(id=report_id).delete()
		return HttpResponseRedirect('/reports/')
	else:
		return render(request, 'reporter/remove.html')


@login_required()
def create(request):
	import updater

	if request.method == 'POST':
		form = ReportForm(request.POST)
		if form.is_valid():
			title = request.POST.get('title_text')
			content = request.POST.get('content_text')
			pub = timezone.now()
			report = Report(title_text=title, content_text=content, pub_date=pub)
			report.save()
			updater.start()
			return HttpResponseRedirect('/reports/' + str(report.id))
	else:
		return render(request, 'reporter/create.html')


@login_required()
def edit(request, report_id):
	report = get_object_or_404(Report, id=report_id)
	if request.method == 'POST':
		report.title_text = request.POST.get('title_text')
		report.content_text = request.POST.get('content_text')
		report.save()
		return HttpResponseRedirect('/reports/')
	else:
		return render(request, 'reporter/edit.html', {'report': report})


@login_required()
def download_excel(request, report_id):
	from openpyxl import Workbook
	from openpyxl.styles import Font, Alignment
	from openpyxl.utils import get_column_letter

	# TODO: Need to paste title and content of report

	file_name = 'Report ' + str(report_id) + '.xlsx'

	wb = Workbook()
	ws = wb.active
	ws.title = file_name

	ws['A1'] = 'ID'
	ws['A1'].font = Font(bold=True)
	ws['A1'].alignment = Alignment(horizontal='center')
	ws['B1'] = 'Name'
	ws['B1'].font = Font(bold=True)
	ws['B1'].alignment = Alignment(horizontal='center')
	ws['C1'] = 'State'
	ws['C1'].font = Font(bold=True)
	ws['C1'].alignment = Alignment(horizontal='center')
	ws['D1'] = 'RAM (kb)'
	ws['D1'].font = Font(bold=True)
	ws['D1'].alignment = Alignment(horizontal='center')
	ws['E1'] = 'Date'
	ws['E1'].font = Font(bold=True)
	ws['E1'].alignment = Alignment(horizontal='center')

	processes = Process.objects.all().filter(report=report_id)

	for i, process in enumerate(processes):
		ws.cell(row=i + 2, column=1, value=processes[i].id)
		ws.cell(row=i + 2, column=2, value=processes[i].name_text)
		ws.cell(row=i + 2, column=3, value=processes[i].state_text)
		ws.cell(row=i + 2, column=4, value=processes[i].ram_in_kb_int)
		ws.cell(row=i + 2, column=5, value=processes[i].date)

	column_widths = []

	for i, cell in enumerate(ws[1]):
		column_widths += [len(cell.value)]

	for row in ws.rows:
		for i, cell in enumerate(row):
			length = len(str(cell.value)) + 2
			if length > column_widths[i]:
				column_widths[i] = length
			else:
				column_widths += [length]

	for i, column_width in enumerate(column_widths):
		ws.column_dimensions[get_column_letter(i + 1)].width = column_width

	wb.save(file_name)
	wb.close()

	f = open(file_name, 'rb')
	response = HttpResponse(f.read())
	f.close()

	file_type = mimetypes.guess_type(file_name)
	if file_type is None:
		file_type = 'application/octet-stream'

	response['Content-Type'] = file_type
	response['Content-Length'] = str(os.stat(file_name).st_size)
	response['Content-Disposition'] = "attachment; filename=%s" % file_name

	os.remove(file_name)

	return response


@login_required()
def download_word(request, report_id):
	import docx

	file_name = 'Report ' + str(report_id) + '.docx'

	report = Report.objects.get(id=report_id)

	doc = docx.Document()

	doc.add_heading("Report " + str(report_id) + " - " + report.title_text, 0)
	doc.add_paragraph("Description: " + report.content_text, 'List Paragraph')
	doc.add_paragraph("Created: " + str(report.pub_date), 'List Paragraph').runs[0].style = 'Subtle Emphasis'

	processes = Process.objects.all().filter(report=report_id)

	for process in processes:
		doc.add_heading("Process " + str(process.id) + " - " + process.name_text, 1)
		doc.add_paragraph("State: ", 'List Bullet').add_run(process.state_text).bold = True
		doc.add_paragraph("RAM (kb): ", 'List Bullet').add_run(str(process.ram_in_kb_int)).bold = True
		doc.add_paragraph("Date: ", 'List Bullet').add_run(str(process.date)).bold = True

	doc.save(file_name)

	f = open(file_name, 'rb')
	response = HttpResponse(f.read())
	f.close()

	file_type = mimetypes.guess_type(file_name)
	if file_type is None:
		file_type = 'application/octet-stream'

	response['Content-Type'] = file_type
	response['Content-Length'] = str(os.stat(file_name).st_size)
	response['Content-Disposition'] = "attachment; filename=%s" % file_name

	os.remove(file_name)

	return response


def user_login(request):
	# TODO: Fix bug with bad redirect after login

	next_page = ""
	if request.method == 'GET':
		next_page = request.GET.get('next')

	if request.method == 'POST':
		username = request.POST['username']
		password = request.POST['password']
		user = authenticate(username=username, password=password)

		if user is not None:
			if user.is_active:
				login(request, user)
				print(next_page)
				if next_page != "":
					return HttpResponseRedirect(next_page)
				else:
					return HttpResponseRedirect('/reports/')
			else:
				message = "User ' + user.username + ' is not active, create new user"
		else:
			message = "Incorrect username or password"

		return render(request, 'reporter/login.html', {'message': message, 'username': username})
	else:
		return render(request, 'reporter/login.html')


def user_logout(request):
	logout(request)
	return HttpResponseRedirect('')
